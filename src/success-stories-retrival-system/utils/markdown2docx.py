from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString
from io import BytesIO
from markdown_it import MarkdownIt
import logging
import requests
import base64
from PIL import Image


def add_hyperlink(paragraph, url, text, color='0563C1', underline=True):
    """
    Add a hyperlink to a paragraph
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        url: The URL to link to
        text: The text to display for the link
        color: The color of the link (hex format without #)
        underline: Whether to underline the link
    
    Returns:
        The hyperlink element
    """
    # Create the hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a new run for the hyperlink text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set the color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    
    # Set underline if needed
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink

def markdown_to_docx(md_content:str , output_filename=None):
    """
    Convert Markdown content to a Word document with table support
    
    Args:
        soup: BeautifulSoup object containing HTML
        output_filename: Name of the output .docx file (if None, returns BytesIO)
    
    Returns:
        BytesIO object if output_filename is None, otherwise saves to file
    """
    
    md = MarkdownIt("gfm-like")
    html_content = md.render(md_content)
    print("-----------------------html content for document:------------------\n",html_content)
    # print("html_content from markdown paresr:\n",html_content)
    soup = BeautifulSoup(html_content, features="lxml")
    doc = Document()
    
    def calculate_image_size(image_stream, max_width=6.0, max_height=8.0):
        """
        Calculate appropriate image size maintaining aspect ratio
        
        Args:
            image_stream: BytesIO stream containing image
            max_width: Maximum width in inches (default 6.0 for standard document)
            max_height: Maximum height in inches (default 8.0)
        
        Returns:
            tuple: (width, height) in Inches, or None if image can't be processed
        """
        try:
            image_stream.seek(0)
            img = Image.open(image_stream)
            img_width, img_height = img.size
            
            # Calculate aspect ratio
            aspect_ratio = img_width / img_height
            
            # Determine best fit
            if aspect_ratio > 1:  # Landscape
                # Width is limiting factor
                width = min(max_width, max_width)
                height = width / aspect_ratio
                if height > max_height:
                    height = max_height
                    width = height * aspect_ratio
            else:  # Portrait or square
                # Height is limiting factor
                height = min(max_height, max_height * 0.6)  # Limit portrait height
                width = height * aspect_ratio
                if width > max_width:
                    width = max_width
                    height = width / aspect_ratio
            
            # Reset stream position
            image_stream.seek(0)
            return Inches(width), Inches(height)
        except Exception as e:
            print(f"Warning: Could not calculate image size: {e}")
            image_stream.seek(0)
            return Inches(5), None  # Return default width only
    
    def process_element(element, parent_container=None, list_level=0):
        """Recursively process HTML elements and add to document"""
        
        if parent_container is None:
            parent_container = doc
            
        if isinstance(element, str):
            # Handle text nodes
            if element.strip():
                if hasattr(parent_container, 'add_paragraph'):
                    p = parent_container.add_paragraph(element.strip())
                    p.style = 'Normal'
            return
        
        # Handle different HTML tags
        tag_name = element.name
        
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Headers
            level = int(tag_name[1])
            heading = doc.add_paragraph(element.get_text().strip())
            heading.style = f'Heading {level}'
            
        elif tag_name == 'p':
            # Paragraphs - check if it contains <br> tags
            if element.find('br'):
                process_paragraph_with_breaks(element, doc)
            else:
                p = doc.add_paragraph()
                process_inline_elements(element, p)
            
        elif tag_name in ['ul', 'ol']:
            # Lists
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(li.get_text().strip())
                if tag_name == 'ul':
                    p.style = 'List Bullet'
                else:
                    p.style = 'List Number'
                    
        elif tag_name == 'table':
            # Tables
            process_table(element, doc)
            
        elif tag_name == 'hr':
            # Horizontal rule
            doc.add_paragraph('─' * 50)
            
        elif tag_name == 'br':
            # Line break - standalone
            doc.add_paragraph()
            
        elif tag_name == 'blockquote':
            # Blockquote
            p = doc.add_paragraph(element.get_text().strip())
            p.style = 'Quote'
            
        elif tag_name == 'pre':
            code_tag = element.find('code')
            if code_tag:
                lang_class = code_tag.get('class', [])
                language = None
                if lang_class:
                    for cls in lang_class:
                        if cls.startswith('language-'):
                            language = cls.replace('language-', '')
                            break
                
                code_text = code_tag.get_text()
                
                # Handle Mermaid diagrams
                if language == 'mermaid':
                    try:
                        # Strip whitespace and encode with base64url
                        code_text_clean = code_text.strip()
                        encoded = base64.urlsafe_b64encode(code_text_clean.encode('utf-8')).decode('ascii')
                        url = f"https://mermaid.ink/img/{encoded}"
                        
                        response = requests.get(url, timeout=30)
                        print("image url:",url)
                        
                        if response.status_code == 200:
                            image_stream = BytesIO(response.content)
                            
                            # Calculate appropriate image size
                            width, height = calculate_image_size(image_stream)
                            
                            # Add image with calculated dimensions
                            if height:
                                doc.add_picture(image_stream, width=width, height=height)
                            else:
                                doc.add_picture(image_stream, width=width)
                            doc.add_paragraph()
                        else:
                            print(f"The reponse from mermaid.ink:{response.content}")
                            raise Exception(f"Status {response.status_code}")
                    except Exception as e:
                        # Fallback: show as code
                        logging.info(f"Error occured while generating mermaid img:{e}")
                        p = doc.add_paragraph("[Mermaid Diagram - Render Failed] ")
                        add_hyperlink(p, url, "Click here to view diagram", color='0563C1', underline=True)
                        p.runs[0].font.name = 'Courier New'
                        p.runs[0].font.size = Pt(9)
                        doc.add_paragraph()
                else:
                    # Regular code block
                    if language:
                        label_p = doc.add_paragraph(f"[{language}]")
                        label_p.runs[0].font.size = Pt(8)
                        label_p.runs[0].italic = True
                    
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    doc.add_paragraph()
                
        elif tag_name == 'code':
            # Inline code (if not already handled by <pre>)
            # Skip if parent is <pre>
            if element.parent and element.parent.name == 'pre':
                pass
            else:
                # Inline code - will be handled by parent paragraph
                pass
                
        elif tag_name in ['strong', 'b', 'em', 'i', 'u', 'a']:
            # Inline elements - will be handled by parent
            pass
        
        else:
            # For other tags, process children
            for child in element.children:
                process_element(child, parent_container, list_level)
    
    def process_paragraph_with_breaks(element, doc):
        """Process paragraph that contains <br> tags by splitting into multiple paragraphs"""
        p = doc.add_paragraph()
        
        for child in element.children:
            if isinstance(child, NavigableString):
                # Text node
                text = str(child).strip()
                if text:
                    run = p.add_run(text)
                    # Apply formatting from parent
                    apply_formatting(run, child.parent, element)
            elif child.name == 'br':
                # Line break - create new paragraph
                p = doc.add_paragraph()
            elif child.name in ['strong', 'b', 'em', 'i', 'u']:
                # Inline formatting
                text = child.get_text()
                if text.strip():
                    run = p.add_run(text)
                    if child.name in ['strong', 'b']:
                        run.bold = True
                    elif child.name in ['em', 'i']:
                        run.italic = True
                    elif child.name == 'u':
                        run.underline = True
            elif child.name == 'a':
                # Link
                href = child.get('href', '#')
                text = child.get_text()
                if text.strip():
                    add_hyperlink(p, href, text)
            else:
                # Recurse for other elements
                text = child.get_text()
                if text.strip():
                    run = p.add_run(text)
    
    def apply_formatting(run, parent, stop_at):
        """Apply formatting from parent elements"""
        current = parent
        while current and current != stop_at:
            if current.name in ['strong', 'b']:
                run.bold = True
            elif current.name in ['em', 'i']:
                run.italic = True
            elif current.name == 'u':
                run.underline = True
            elif current.name == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            current = current.parent
    
    def process_inline_elements(element, paragraph):
        """Process inline elements like bold, italic, links within a paragraph"""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    run = paragraph.add_run(text)
            elif child.name == 'br':
                # Add line break within paragraph
                paragraph.add_run().add_break()
            elif child.name in ['strong', 'b', 'em', 'i', 'u']:
                text = child.get_text().strip()
                if text:
                    run = paragraph.add_run(text)
                    if child.name in ['strong', 'b']:
                        run.bold = True
                    elif child.name in ['em', 'i']:
                        run.italic = True
                    elif child.name == 'u':
                        run.underline = True
            elif child.name == 'a':
                href = child.get('href', '#')
                text = child.get_text()
                if text.strip():
                    add_hyperlink(paragraph, href, text)
            elif child.name == 'code':
                text = child.get_text()
                if text:
                    run = paragraph.add_run(text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
    
    def process_table(table_element, doc):
        """Process HTML table and add to document"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Determine number of columns
        max_cols = 0
        for row in rows:
            cols = len(row.find_all(['td', 'th']))
            max_cols = max(max_cols, cols)
        
        if max_cols == 0:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        
        # Populate table
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    table_cell = table.rows[row_idx].cells[col_idx]
                    # Get cell text
                    cell_text = cell.get_text().strip()
                    table_cell.text = cell_text
                    
                    # Format header cells
                    if cell.name == 'th':
                        for paragraph in table_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(11)
        
        # Add spacing after table
        doc.add_paragraph()
    
    # Process all top-level elements
    for element in soup.children:
        process_element(element)
    
    # Return BytesIO or save to file
    if output_filename is None:
        # Return as BytesIO
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)  # Reset pointer to beginning
        return byte_io
    else:
        # Save to file
        doc.save(output_filename)
        print(f"Document saved as '{output_filename}'")
        return doc